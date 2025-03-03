import fitz
from docx import Document

ALLOWED_EXTENSIONS = {"pdf", "docx"}

# Allow only PDFS and Word Documents for upload
def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# Extract text from PDF
def extract_text_from_pdf(filepath):
    doc = fitz.open(filepath)
    return "\n".join(page.get_text() for page in doc)

# Extract text from PDF directly from file object
def extract_text_from_pdf_stream(file):
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        text = "\n".join([page.get_text() for page in doc])
    return text

# Extract text from doc file
def extract_text_from_docx(filepath):
    doc = Document(filepath)
    return "\n".join([para.text for para in doc.paragraphs])


# Extract text from doc file directly from file object
def extract_text_from_docx_stream(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

